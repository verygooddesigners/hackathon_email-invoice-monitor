"""Word document parser — extracts text from .docx files using python-docx."""

import io
import logging
from docx import Document

logger = logging.getLogger("invoice_monitor")


def extract_docx_text(source) -> str:
    """Extract all text from a Word document (.docx).

    Extracts text from paragraphs and table cells, since invoices are often
    formatted as tables.

    Args:
        source: File path (str) or bytes object of the Word document.

    Returns:
        Full text content of the document, or an error message string.
    """
    try:
        if isinstance(source, bytes):
            doc = Document(io.BytesIO(source))
        elif isinstance(source, str):
            if not source.lower().endswith((".docx",)):
                return "[ERROR] Unsupported file format. Only .docx files are supported (not .doc)."
            doc = Document(source)
        else:
            return "[ERROR] Invalid input type for docx parser."

        text_parts = []

        # Extract paragraph text
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        # Extract table cell text (critical for invoice line items)
        for table in doc.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    text_parts.append(" | ".join(row_cells))

        return "\n".join(text_parts) if text_parts else ""

    except Exception as e:
        logger.error(f"Error parsing Word document: {e}")
        return f"[ERROR] Could not parse Word document: {e}"
